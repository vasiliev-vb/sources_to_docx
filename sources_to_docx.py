#!/usr/bin/env python3
import argparse
import docx
import os

def modify_document(d, file_lists, args):
    codep = None
    for paragraph in d.paragraphs:
        if args.code_place_mark in paragraph.text:
            codep = paragraph
    if codep is not None:
        add_paragraph = codep.insert_paragraph_before
    else:
        add_paragraph = d.add_paragraph
        print("Paragraph with '{}' text was not found. Appending to end of file".format(args.code_place_mark))

    p = None
    for file_list in file_lists:
        ## insert page break after each paragraph except last one
        if args.heading_page_break and p is not None:
            run = p.add_run()
            run.add_break(docx.enum.text.WD_BREAK.PAGE)

        if args.with_headings:
            p = add_paragraph(args.file_list_description.format(file_list = os.path.splitext(os.path.basename(file_list))[0]))
            p.style = args.heading_style
        with open(file_list,'r') as file_list_content:
            for file_name in file_list_content:
                file_name = file_name.strip()
                file_name_base = os.path.basename(file_name)
                file_name_noext = os.path.splitext(file_name_base)[0]
                p = add_paragraph(args.file_description.format(file_name = file_name_base, file_name_noext = file_name_noext))
                p.style = args.text_style
                with open(file_name,'r') as file_content:
                    if args.verbose:
                        print('Inserting {}'.format(file_name))
                    p = add_paragraph(file_content.read())
                    p.style = args.code_style
    if codep:
        codep.text=''
    return d

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Writes a bunch of text files to .docx')

    parser.add_argument('--input','-i',            required=True, nargs='+', help='Input file list or multiple file lists. In latter case, hierarchy can be used')
    parser.add_argument('--template','-t',         default=None,             help='Template .docx file with styles and preformatting')
    parser.add_argument('--output', '-o',          required=True,            help='Output .docx file')

    parser.add_argument('--code_place_mark',       default='<CODE>',         help='Paragraph with this mark will be replaced by code')
                                                                             
    parser.add_argument('--file_list_description', default='{file_list}',    help='Description of file list')
    parser.add_argument('--file_description',      default='{file_name}:',   help='Description of file')
                                                                             
    parser.add_argument('--with_headings',         action='store_true',      help='Print headings for files')
    parser.add_argument('--heading_style',         default='Heading 1',      help='Style for text')
    parser.add_argument('--heading_page_break',    action='store_true',      help='Add page break between headings')
                                                                             
    parser.add_argument('--text_style',            default='Normal',         help='Style for text')
    parser.add_argument('--code_style',            default='Normal',         help='Style for code')
                                                                             
    parser.add_argument('--print_styles',          action='store_true',      help='Print styles available in template')
    parser.add_argument('--verbose','-v',          action='store_true',      help='Print various information')

    args = parser.parse_args()

    if args.template is not None:
        d = docx.Document(args.template)
    else:
        d = docx.Document()

    if args.print_styles:
        print('=================')
        print('Styles available:')
        for s in d.styles:
            print(s.name) 
        print('=================')
        exit(0)

    d = modify_document(d, args.input, args)
    d.save(args.output)

